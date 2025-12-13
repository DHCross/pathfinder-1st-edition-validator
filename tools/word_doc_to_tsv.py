#!/usr/bin/env python3
"""
Word Document Table Extractor
=============================

Extracts ALL tables from a Microsoft Word (.docx) document and converts them
to tab-delimited format for InDesign import.

This is the "holy grail" tool - point it at a Word document and it extracts
every table, converting them to clean TSV format.

Usage:
    python3 tools/word_doc_to_tsv.py document.docx -o output.txt
    python3 tools/word_doc_to_tsv.py document.docx --clipboard
    python3 tools/word_doc_to_tsv.py document.docx --per-table  # One file per table

Requirements:
    pip install python-docx

Author: Pathfinder 1st Edition Validator Tools
Date: 2025-12-13
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_cell_text(cell) -> str:
    """
    Extract text from a Word table cell, handling nested paragraphs.
    
    Args:
        cell: A docx table cell object
    
    Returns:
        Cleaned text content
    """
    # Get all paragraphs in the cell
    paragraphs = [p.text for p in cell.paragraphs]
    text = ' '.join(paragraphs)
    
    # Clean up whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


def extract_table(table: 'Table', table_index: int) -> Tuple[str, List[List[str]]]:
    """
    Extract a single table from a Word document.
    
    Args:
        table: A docx Table object
        table_index: Index of the table (for naming)
    
    Returns:
        Tuple of (header_name, rows)
    """
    rows = []
    
    for row in table.rows:
        cells = [extract_cell_text(cell) for cell in row.cells]
        
        # Handle merged cells - docx repeats content for merged cells
        # We'll deduplicate adjacent identical cells
        deduped_cells = []
        prev_cell = None
        for cell in cells:
            if cell != prev_cell:
                deduped_cells.append(cell)
                prev_cell = cell
            else:
                # Keep empty placeholder for merged cells
                deduped_cells.append('')
        
        rows.append(deduped_cells)
    
    # Try to detect a header from the first row
    header = f"Table {table_index + 1}"
    
    return header, rows


def extract_tables_from_docx(docx_path: Path) -> List[Tuple[str, List[List[str]]]]:
    """
    Extract all tables from a Word document.
    
    Args:
        docx_path: Path to the .docx file
    
    Returns:
        List of (header, rows) tuples
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed. Install with: pip install python-docx", 
              file=sys.stderr)
        sys.exit(1)
    
    doc = Document(docx_path)
    tables = []
    
    for i, table in enumerate(doc.tables):
        header, rows = extract_table(table, i)
        if rows:  # Only include non-empty tables
            tables.append((header, rows))
    
    return tables


def format_as_tsv(tables: List[Tuple[str, List[List[str]]]], 
                  include_headers: bool = True,
                  normalize_columns: bool = True) -> str:
    """
    Format extracted tables as tab-separated values.
    
    Args:
        tables: List of (header, rows) tuples
        include_headers: Whether to include section headers
        normalize_columns: Pad rows to have consistent column count
    
    Returns:
        Tab-delimited string
    """
    output_lines = []
    
    for header, rows in tables:
        if not rows:
            continue
            
        # Normalize column count if requested
        if normalize_columns and rows:
            max_cols = max(len(row) for row in rows)
            rows = [row + [''] * (max_cols - len(row)) for row in rows]
        
        if include_headers and header:
            output_lines.append(f"# {header}")
            output_lines.append("")
        
        for row in rows:
            output_lines.append('\t'.join(row))
        
        output_lines.append("")
    
    return '\n'.join(output_lines)


def format_as_markdown(tables: List[Tuple[str, List[List[str]]]]) -> str:
    """
    Format extracted tables as Markdown pipe tables.
    
    Args:
        tables: List of (header, rows) tuples
    
    Returns:
        Markdown string with pipe tables
    """
    output_lines = []
    
    for header, rows in tables:
        if not rows:
            continue
        
        # Normalize column count
        max_cols = max(len(row) for row in rows)
        rows = [row + [''] * (max_cols - len(row)) for row in rows]
        
        # Add header
        output_lines.append(f"## {header}")
        output_lines.append("")
        
        # First row as header
        if rows:
            header_row = rows[0]
            output_lines.append("| " + " | ".join(header_row) + " |")
            output_lines.append("| " + " | ".join(["---"] * len(header_row)) + " |")
            
            # Data rows
            for row in rows[1:]:
                output_lines.append("| " + " | ".join(row) + " |")
        
        output_lines.append("")
    
    return '\n'.join(output_lines)


def convert_docx(input_path: Path, output_path: Optional[Path] = None,
                 include_headers: bool = True, clipboard: bool = False,
                 per_table: bool = False, output_format: str = 'tsv') -> str:
    """
    Convert Word document tables to tab-delimited or Markdown format.
    
    Args:
        input_path: Path to input .docx file
        output_path: Path to output file (optional)
        include_headers: Include section headers
        clipboard: Copy to clipboard instead of saving
        per_table: Save each table to a separate file
        output_format: 'tsv' or 'markdown'
    
    Returns:
        The converted content
    """
    tables = extract_tables_from_docx(input_path)
    
    if not tables:
        print("Warning: No tables found in document", file=sys.stderr)
        return ""
    
    print(f"Found {len(tables)} table(s) in document")
    
    # Format output
    if output_format == 'markdown':
        content = format_as_markdown(tables)
        ext = '.md'
    else:
        content = format_as_tsv(tables, include_headers)
        ext = '.txt'
    
    # Handle per-table output
    if per_table and output_path:
        base = output_path.stem
        parent = output_path.parent
        
        for i, (header, rows) in enumerate(tables):
            table_content = format_as_tsv([(header, rows)], include_headers) if output_format == 'tsv' \
                           else format_as_markdown([(header, rows)])
            
            table_path = parent / f"{base}_table{i + 1}{ext}"
            with open(table_path, 'w', encoding='utf-8') as f:
                f.write(table_content)
            print(f"✓ Saved: {table_path}")
        
        return content
    
    # Output
    if clipboard:
        try:
            import pyperclip
            pyperclip.copy(content)
            print("✓ Copied to clipboard!")
        except ImportError:
            print("Error: pyperclip not installed. Install with: pip install pyperclip", 
                  file=sys.stderr)
            print("\nContent:")
            print(content)
    elif output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✓ Saved to: {output_path}")
    else:
        print(content)
    
    return content


def main():
    parser = argparse.ArgumentParser(
        description='Extract ALL tables from a Word document to tab-delimited or Markdown format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Extract all tables to single TSV file
  %(prog)s document.docx -o tables.txt
  
  # Extract to Markdown format
  %(prog)s document.docx -o tables.md --format markdown
  
  # One file per table
  %(prog)s document.docx -o tables.txt --per-table
  # Creates: tables_table1.txt, tables_table2.txt, etc.
  
  # Copy to clipboard
  %(prog)s document.docx --clipboard
  
  # Print to stdout
  %(prog)s document.docx
  
  # Without section headers
  %(prog)s document.docx --no-headers -o output.txt

Requirements:
  pip install python-docx
        """
    )
    
    parser.add_argument('input', type=Path, help='Input Word document (.docx)')
    parser.add_argument('-o', '--output', type=Path, help='Output file path')
    parser.add_argument('--no-headers', action='store_true',
                       help='Exclude section headers from output')
    parser.add_argument('--clipboard', action='store_true',
                       help='Copy result to clipboard (requires pyperclip)')
    parser.add_argument('--per-table', action='store_true',
                       help='Save each table to a separate file')
    parser.add_argument('--format', choices=['tsv', 'markdown'], default='tsv',
                       help='Output format (default: tsv)')
    
    args = parser.parse_args()
    
    # Check for python-docx
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed.", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        return 1
    
    # Validate input
    if not args.input.exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        return 1
    
    if not args.input.suffix.lower() == '.docx':
        print(f"Warning: File does not have .docx extension: {args.input}", file=sys.stderr)
    
    try:
        convert_docx(
            args.input,
            args.output,
            include_headers=not args.no_headers,
            clipboard=args.clipboard,
            per_table=args.per_table,
            output_format=args.format
        )
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
